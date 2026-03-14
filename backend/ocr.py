"""
ocr.py — Text extraction pipeline for ContextBridge
=====================================================
Extraction strategy by file type:

  .txt / .md / .csv   → decode bytes directly
  .docx               → python-docx (paragraphs + tables)
  .pdf                → 1) pdfplumber (native text)
                        2) PyMuPDF → page images → Gemini Vision (scanned PDFs)
  image files         → Gemini Vision (multimodal image-to-text)
  anything else       → attempt UTF-8 decode, raise on failure

Uses Google Gemini's multimodal vision capability for image text extraction.
"""

from __future__ import annotations

import io
import os
from pathlib import Path
from typing import Tuple

import google.generativeai as genai
from PIL import Image
from dotenv import load_dotenv

load_dotenv()

# ── Configure Gemini ─────────────────────────────────────────────────────────
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)

GEMINI_MODEL = "gemini-2.5-flash"

# ── Supported image extensions ────────────────────────────────────────────────
_IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".tif", ".webp"}

# ── Minimum character threshold to accept native PDF extraction ───────────────
_NATIVE_PDF_MIN_CHARS = 80

# ── Max PDF pages to OCR (avoids very long processing times) ─────────────────
_MAX_PDF_PAGES = 25

# ── PyMuPDF render zoom (higher = better OCR quality, more memory) ────────────
_PDF_RENDER_ZOOM = 2.0

# ── OCR prompt for Gemini Vision ─────────────────────────────────────────────
_OCR_PROMPT = """You are an OCR assistant. Extract ALL text visible in this image.
Rules:
- Return ONLY the extracted text, nothing else.
- Preserve the original formatting, layout and line breaks as closely as possible.
- Include all text: headings, body text, labels, captions, numbers, dates, etc.
- If the image contains a table, format it with pipes (|) and dashes (-).
- If no text is visible, return exactly: [NO TEXT DETECTED]
- Do NOT add any commentary, explanation, or markdown formatting."""


# ─────────────────────────────────────────────────────────────────────────────
# Internal helpers
# ─────────────────────────────────────────────────────────────────────────────


def _ocr_image_bytes(image_bytes: bytes) -> str:
    """
    Send image bytes to Gemini Vision for text extraction (OCR).

    Uses the Gemini multimodal API to extract all visible text from an image.
    Returns the extracted text string (may be empty).
    """
    image = Image.open(io.BytesIO(image_bytes)).convert("RGB")

    model = genai.GenerativeModel(GEMINI_MODEL)
    response = model.generate_content([_OCR_PROMPT, image])

    text = response.text.strip()

    # If Gemini says no text, return empty
    if text == "[NO TEXT DETECTED]":
        return ""

    return text


def _extract_pdf_native(file_bytes: bytes) -> str:
    """
    Attempt to extract native (selectable) text from a PDF using pdfplumber.
    Returns an empty string if extraction fails or yields nothing useful.
    """
    try:
        import pdfplumber  # type: ignore[import]
    except ImportError:
        return ""

    try:
        text_parts: list[str] = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n".join(text_parts).strip()
    except Exception as exc:  # noqa: BLE001
        print(f"[OCR] pdfplumber native extraction failed: {exc}")
        return ""


def _extract_pdf_via_vision(file_bytes: bytes) -> str:
    """
    Render each PDF page to a PNG image using PyMuPDF, then extract text
    with Gemini Vision. Used as a fallback for scanned / image-only PDFs.
    """
    try:
        import fitz  # type: ignore[import]  # PyMuPDF
    except ImportError as exc:
        raise ImportError("PyMuPDF is not installed. Run: pip install PyMuPDF") from exc

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    total_pages = min(doc.page_count, _MAX_PDF_PAGES)
    text_parts: list[str] = []

    try:
        mat = fitz.Matrix(_PDF_RENDER_ZOOM, _PDF_RENDER_ZOOM)
        for page_num in range(total_pages):
            page = doc[page_num]
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img_bytes = pix.tobytes("png")

            try:
                page_text = _ocr_image_bytes(img_bytes)
                if page_text:
                    text_parts.append(page_text)
                    print(f"[OCR] Page {page_num + 1}/{total_pages} extracted via Gemini Vision.")
            except Exception as exc:  # noqa: BLE001
                print(f"[OCR] Gemini Vision failed on PDF page {page_num + 1}: {exc}")
    finally:
        doc.close()

    return "\n\n".join(text_parts).strip()


def _extract_docx(file_bytes: bytes) -> str:
    """
    Extract text from a .docx file using python-docx.
    Includes paragraphs and table cell content.
    """
    try:
        from docx import Document as DocxDocument  # type: ignore[import]
    except ImportError as exc:
        raise ImportError(
            "python-docx is not installed. Run: pip install python-docx"
        ) from exc

    doc = DocxDocument(io.BytesIO(file_bytes))
    parts: list[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        stripped = para.text.strip()
        if stripped:
            parts.append(stripped)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                parts.append(" | ".join(row_cells))

    return "\n".join(parts).strip()


# ─────────────────────────────────────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────────────────────────────────────


def extract_text(file_bytes: bytes, filename: str) -> Tuple[str, str]:
    """
    Extract plain text from a file.

    Parameters
    ----------
    file_bytes : bytes
        Raw content of the uploaded file.
    filename : str
        Original filename (used to determine the extraction strategy).

    Returns
    -------
    (text, method) : Tuple[str, str]
        text   — Extracted plain text (may be empty on total failure).
        method — One of: 'text' | 'native' | 'vision' | 'failed'

    Raises
    ------
    ValueError
        If the file type is completely unsupported and cannot even be
        decoded as plain text.
    """
    suffix = Path(filename).suffix.lower()

    # ── Plain text ────────────────────────────────────────────────────────────
    if suffix in (".txt", ".md", ".csv", ".json", ".xml", ".html", ".htm"):
        try:
            return file_bytes.decode("utf-8", errors="replace").strip(), "text"
        except Exception:
            return file_bytes.decode("latin-1", errors="replace").strip(), "text"

    # ── DOCX ─────────────────────────────────────────────────────────────────
    if suffix == ".docx":
        try:
            return _extract_docx(file_bytes), "native"
        except Exception as exc:
            raise ValueError(f"Could not extract text from DOCX: {exc}") from exc

    # ── PDF ───────────────────────────────────────────────────────────────────
    if suffix == ".pdf":
        # 1) Try native text extraction first (fast, free)
        native_text = _extract_pdf_native(file_bytes)
        if len(native_text) >= _NATIVE_PDF_MIN_CHARS:
            return native_text, "native"

        # 2) Scanned / image-only PDF → Gemini Vision OCR
        print(
            f"[OCR] '{filename}' native text too short "
            f"({len(native_text)} chars), switching to Gemini Vision…"
        )
        try:
            vision_text = _extract_pdf_via_vision(file_bytes)
            if vision_text:
                return vision_text, "vision"
            # Both paths yielded nothing
            return native_text or "", "failed"
        except Exception as exc:
            print(f"[OCR] Gemini Vision failed for '{filename}': {exc}")
            return native_text or "", "failed"

    # ── Images ────────────────────────────────────────────────────────────────
    if suffix in _IMAGE_EXTS:
        try:
            text = _ocr_image_bytes(file_bytes)
            return text, "vision"
        except Exception as exc:
            raise ValueError(
                f"Vision OCR failed for image '{filename}': {exc}"
            ) from exc

    # ── Fallback: try plain-text decode ───────────────────────────────────────
    try:
        decoded = file_bytes.decode("utf-8", errors="replace").strip()
        if decoded:
            return decoded, "text"
    except Exception:
        pass

    raise ValueError(
        f"Unsupported file type '{suffix}'. "
        "Supported: PDF, DOCX, TXT, MD, CSV, JPG, PNG, TIFF, WEBP, BMP, GIF."
    )
